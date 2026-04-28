"""
Script de geração de avaliação impressa no padrão Theodoro de Bona.
Uso: python3 generate.py <arquivo_json_questoes> <nome_disciplina>
O JSON deve ter: [{"enunciado": "...", "alternativas": ["a) ...", "b) ...", ...]}, ...]
"""
import json, sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE = "/root/.hermes/cache/documents/doc_81bd7d337c3d_Avaliacao_Bona_Template.docx"
LETTERS = ['a', 'b', 'c', 'd']

def generate(questions_file, output_name):
    with open(questions_file) as f:
        questions = json.load(f)
    
    doc = Document(TEMPLATE)
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Instruções
    inst = doc.add_paragraph()
    run = inst.add_run("Instruções: "); run.bold = True; run.font.size = Pt(10)
    run = inst.add_run("Leia atentamente cada questão e assinale apenas uma alternativa correta. "
                        "Não rasure. Cada questão vale 0,6 pontos.")
    run.font.size = Pt(10)
    doc.add_paragraph()
    
    # Questões
    for q in questions:
        qp = doc.add_paragraph()
        qp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = qp.add_run(q["enunciado"]); run.bold = True; run.font.size = Pt(11)
        for i, alt in enumerate(q["alternativas"]):
            ap = doc.add_paragraph()
            ap.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ap.paragraph_format.left_indent = Cm(1.2)
            run = ap.add_run(f"{LETTERS[i]}) {alt}"); run.font.size = Pt(10.5)
    
    # Rodapé
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("BOA PROVA!"); run.bold = True; run.font.size = Pt(12)
    
    output = f"/tmp/{output_name}.docx"
    doc.save(output)
    print(f"✅ {output}")
    return output

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Uso: python3 generate.py <questions.json> <output_name>")
        sys.exit(1)
    generate(sys.argv[1], sys.argv[2])
